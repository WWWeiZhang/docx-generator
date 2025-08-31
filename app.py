import streamlit as st
import os
from tempfile import TemporaryDirectory
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl
import zipfile

# ======================== 图片替换规则 ========================
REPLACEMENT_RULES = [
    {"type": "single", "placeholder": "{{IMAGE_1}}", "image_index": 0, "width": Inches(4.5)},
    {"type": "double", "placeholder": "{{IMAGES_2_3}}", "image_indices": [1,2], "width": Inches(2.5)},
    {"type": "single", "placeholder": "{{IMAGE_4}}", "image_index": 3, "width": Inches(4)}
]

# ======================== 图片替换函数 ========================
def replace_single_image(doc, placeholder, path, width):
    for p in doc.paragraphs:
        if placeholder in p.text:
            p.clear()
            p.add_run().add_picture(path, width=width)
            return True
    return False

def replace_two_images_table(doc, placeholder, path1, path2, width):
    for p in doc.paragraphs:
        if placeholder in p.text:
            _insert_image_table_and_remove_paragraph(p, path1, path2, width)
            return True
    return False

def _insert_image_table_and_remove_paragraph(p, path1, path2, width):
    table = p.part.document.add_table(rows=1, cols=2)
    p._p.addprevious(table._tbl)
    _remove_table_borders(table)
    cell1, cell2 = table.cell(0,0), table.cell(0,1)
    cell1.paragraphs[0].paragraph_format.space_before = Pt(0)
    cell1.paragraphs[0].paragraph_format.space_after = Pt(0)
    cell1.paragraphs[0].add_run().add_picture(path1, width=width)
    cell2.paragraphs[0].paragraph_format.space_before = Pt(0)
    cell2.paragraphs[0].paragraph_format.space_after = Pt(0)
    cell2.paragraphs[0].add_run().add_picture(path2, width=width)
    parent = p._p.getparent()
    parent.remove(p._p)

def _remove_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for b in ['top','left','bottom','right','insideH','insideV']:
                border = OxmlElement(f'w:{b}')
                border.set(qn('w:val'), 'nil')
                tcBorders.append(border)
            tcPr.append(tcBorders)

# ======================== 核心处理函数 ========================
def process_documents(excel_path, template_path, rules, images_path, output_folder):
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    workbook = openpyxl.load_workbook(excel_path)
    sheet = workbook.active

    for row in sheet.iter_rows(min_row=2, values_only=True):
        if not row[0]:
            continue
        doc_basename = str(row[0])
        replacements = {}
        for i in range(1, len(row), 2):
            old_text = row[i]
            new_text = row[i+1] if i+1 < len(row) and row[i+1] else ""
            if old_text:
                replacements[str(old_text)] = str(new_text)

        doc = Document(template_path)

        # 文本替换
        for p in doc.paragraphs:
            for old_text, new_text in replacements.items():
                if old_text in p.text:
                    p.text = p.text.replace(old_text, new_text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for old_text, new_text in replacements.items():
                        if old_text in cell.text:
                            cell.text = cell.text.replace(old_text, new_text)

        # 图片替换
        image_files = sorted([f for f in os.listdir(images_path) if f.lower().endswith(('.png','.jpg','.jpeg'))])
        image_paths = [os.path.join(images_path, f) for f in image_files]

        if image_paths:
            max_index_needed = max([rule.get('image_index', -1) for rule in rules if rule['type']=='single'] +
                                   [idx for rule in rules if rule['type']=='double' for idx in rule.get('image_indices', [])])
            if len(image_paths) > max_index_needed:
                for rule in rules:
                    placeholder = rule['placeholder']
                    if rule['type']=='single':
                        replace_single_image(doc, placeholder, image_paths[rule['image_index']], rule['width'])
                    elif rule['type']=='double':
                        idx1, idx2 = rule['image_indices']
                        replace_two_images_table(doc, placeholder, image_paths[idx1], image_paths[idx2], rule['width'])

        # 保存文档
        doc.save(os.path.join(output_folder, f"{doc_basename}.docx"))

# ======================== Streamlit 界面 ========================
st.title("📄 批量 Word 文档生成工具")

excel_file = st.file_uploader("上传 Excel 文件", type=["xlsx"])
template_file = st.file_uploader("上传 Word 模板", type=["docx"])
images_files = st.file_uploader("上传图片（可多选）", type=["png","jpg","jpeg"], accept_multiple_files=True)

if st.button("开始生成"):
    if not excel_file or not template_file:
        st.error("请上传 Excel 和 Word 模板")
    else:
        with TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            excel_path = tmpdir_path / "data.xlsx"
            template_path = tmpdir_path / "template.docx"
            images_path = tmpdir_path / "images"
            images_path.mkdir()
            output_folder = tmpdir_path / "output"

            # 保存上传文件
            with open(excel_path, "wb") as f:
                f.write(excel_file.read())
            with open(template_path, "wb") as f:
                f.write(template_file.read())
            for f in images_files:
                with open(images_path / f.name, "wb") as file:
                    file.write(f.read())

            # 调用处理函数
            process_documents(str(excel_path), str(template_path), REPLACEMENT_RULES, str(images_path), str(output_folder))

            # 打包 ZIP
            zip_path = tmpdir_path / "result.zip"
            with zipfile.ZipFile(zip_path, 'w') as zipf:
                for file_name in os.listdir(output_folder):
                    zipf.write(output_folder / file_name, file_name)

            with open(zip_path, "rb") as f:
                st.download_button("📥 下载生成的文档", f, file_name="result.zip")
