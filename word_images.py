import os
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl
# ===================================================================
# 1. 图片替换规则配置 (按需编辑)
# ===================================================================
REPLACEMENT_RULES = [
    {
        "type": "single",
        "placeholder": "{{IMAGE_1}}",
        "image_index": 0,
        "width": Inches(4.5)
    },
    {
        "type": "double",
        "placeholder": "{{IMAGES_2_3}}",
        "image_indices": [1, 2],
        "width": Inches(2.5)
    },
    {
        "type": "single",
        "placeholder": "{{IMAGE_4}}",
        "image_index": 3,
        "width": Inches(4)
    }
]


# ===================================================================
# 2. 核心工具函数
# ===================================================================

def replace_single_image(doc, placeholder, path, width):
    for p in doc.paragraphs:
        if placeholder in p.text:
            p.clear()
            p.add_run().add_picture(path, width=width)
            return True
    return False


def replace_two_images_table(doc, placeholder, path1, path2, width):
    """查找并用一个无边框表格替换占位符，并在其中并排放置两张图片。"""
    for p in doc.paragraphs:
        if placeholder in p.text:
            _insert_image_table_and_remove_paragraph(p, path1, path2, width)
            return True
    return False


def _insert_image_table_and_remove_paragraph(p, path1, path2, width):
    """辅助函数：插入一个 1x2 的无边框表格并移除原始段落。"""
    table = p.part.document.add_table(rows=1, cols=2)
    p._p.addprevious(table._tbl)
    _remove_table_borders(table)
    cell1, cell2 = table.cell(0, 0), table.cell(0, 1)
    p1 = cell1.paragraphs[0]
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(0)
    p1.add_run().add_picture(path1, width=width)
    p2 = cell2.paragraphs[0]
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    p2.add_run().add_picture(path2, width=width)
    parent = p._p.getparent()
    parent.remove(p._p)


def _remove_table_borders(table):
    """辅助函数：移除表格的所有边框。"""
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for b in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{b}')
                border.set(qn('w:val'), 'nil')
                tcBorders.append(border)
            tcPr.append(tcBorders)


# ===================================================================
# 3. 主执行逻辑 - 组合了文本和图片替换
# ===================================================================

def process_documents(excel_path, template_path, rules, images_path, output_folder):
    # 确保输出文件夹存在
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
        print(f"已创建输出文件夹: {output_folder}")

    word = None
    try:
        workbook = openpyxl.load_workbook(excel_path)
        sheet = workbook.active

        # 使用python-docx直接处理，无需win32com

        for row in sheet.iter_rows(min_row=2, values_only=True):
            if not row[0]:  # 如果第一列为空，则跳过该行
                continue

            doc_basename = str(row[0])  # 新文件名在Excel第一列

            # 从Excel的第2列、第3列...获取文本替换规则
            replacements = {}
            for i in range(1, len(row), 2):
                old_text = row[i]
                new_text = row[i + 1] if i + 1 < len(row) and row[i + 1] is not None else ""
                if old_text:
                    replacements[str(old_text)] = str(new_text)

            print(f"\n--- 正在处理文档: {doc_basename}.docx ---")

            doc = None
            try:
                # 使用python-docx打开模板
                doc = Document(template_path)

                # --- 文本替换部分 ---
                for p in doc.paragraphs:
                    for old_text, new_text in replacements.items():
                        if old_text in p.text:
                            p.text = p.text.replace(old_text, new_text)

                for table in doc.tables:
                    for cell in table.iter_cells():
                        for old_text, new_text in replacements.items():
                            if old_text in cell.text:
                                cell.text = cell.text.replace(old_text, new_text)

                print(f"  -> 完成文本替换。")

                # --- 图片替换部分 ---
                image_subfolder = os.path.join(images_path, doc_basename)
                if not os.path.isdir(image_subfolder):
                    print(f"  -> 警告: 图片文件夹 '{doc_basename}' 未找到。跳过图片替换。")
                else:
                    image_filenames = sorted(
                        [f for f in os.listdir(image_subfolder) if f.lower().endswith(('.png', '.jpg', '.jpeg'))]
                    )
                    image_paths = [os.path.join(image_subfolder, fname) for fname in image_filenames]

                    max_index_needed = max([rule.get('image_index', -1) for rule in rules if rule['type'] == 'single'] +
                                           [idx for rule in rules if rule['type'] == 'double' for idx in
                                            rule.get('image_indices', [])])

                    if len(image_paths) <= max_index_needed:
                        print(
                            f"  -> 警告: 规则要求 {max_index_needed + 1} 张图片，但只找到 {len(image_paths)} 张。跳过图片替换。")
                    else:
                        for rule in rules:
                            placeholder = rule["placeholder"]
                            print(f"  -> 正在应用图片规则: 查找 '{placeholder}'...")
                            found = False
                            if rule["type"] == "single":
                                img_path = image_paths[rule["image_index"]]
                                found = replace_single_image(doc, placeholder, img_path, rule["width"])
                            elif rule["type"] == "double":
                                idx1, idx2 = rule["image_indices"]
                                img1_path, img2_path = image_paths[idx1], image_paths[idx2]
                                found = replace_two_images_table(doc, placeholder, img1_path, img2_path, rule["width"])
                            if not found:
                                print(f"  -> 警告: 占位符 '{placeholder}' 未在文档中找到。")

                # 所有替换完成后，另存文档
                new_file_path = os.path.join(output_folder, f"{doc_basename}.docx")
                doc.save(new_file_path)  # 使用python-docx的save方法
                print(f"  -> 成功保存至: {new_file_path}")

            except Exception as e:
                print(f"  -> 处理文档'{doc_basename}.docx'时发生错误: {e}")
            finally:
                if doc:
                    del doc

    except Exception as e:
        print(f"启动Word或Excel读取过程中发生严重错误: {e}")
    finally:
        print("\n所有任务处理完毕！")


# ===================================================================
# 4. 运行脚本
# ===================================================================

if __name__ == "__main__":
    # --- 请在这里配置您的文件路径 ---
    excel_file_path = r"D:\Users\zhangwei\Desktop\eraquake\data.xlsx"
    template_word_path = r"D:\Users\zhangwei\Desktop\eraquake\template.docx"
    output_directory = r"D:\Users\zhangwei\Desktop\eraquake\output"
    images_path = r"D:\Users\zhangwei\Desktop\eraquake\images"

    # 运行主函数
    process_documents(excel_file_path, template_word_path, REPLACEMENT_RULES, images_path, output_directory)
